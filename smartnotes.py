#importing modules
from tkinter import *
from tkinter import filedialog
from docx import Document
import os
import tkinter.font as tkFont
from urllib.request import urlopen
from tkdocviewer import *
import splitter
from PIL import Image, ImageTk, ImageFilter
from pdf2image import convert_from_path, convert_from_bytes
import pytesseract
import nltk
from nltk.cluster.util import cosine_distance
import numpy as np
import networkx as nx
from fpdf import FPDF
from docx2pdf import convert
try:
  from nltk.corpus import stopwords
  from nltk.tokenize import sent_tokenize
except:
  nltk.download('stopwords')
  nltk.download('sent_tokenize')
  from nltk.corpus import stopwords
from selenium import webdriver

#opening homepage
root = Tk()
root.title('Sapience')

#title
appTitleImage = Image.open(r'C:\Users\owner\Desktop\Visual Studio Code Projects\SapienceLogo(1).png')
appTitleImage = appTitleImage.resize((506, 156))
appTitleImage = ImageTk.PhotoImage(appTitleImage)
appTitle = Label(root, image=appTitleImage)
appTitle.pack()

#slogan
appSloganFont = tkFont.Font(family='Impact', size=20, slant=tkFont.ITALIC)
appSlogan = Label(root, text='Study Better, Study Smarter', fg='#4670B6', font=appSloganFont)
appSlogan.pack()

#spacer
ButtonFont = tkFont.Font(family='Arial', size=15, weight=tkFont.BOLD)
spacer = Label(root, text='', font=ButtonFont)
spacer.pack()

#SmartNotes
def read_article(file_name):
    file = open(file_name, "r")
    filedata = file.readlines()
    article = filedata[0].split(". ")
    sentences = []

    for sentence in article:
        sentences.append(sentence.replace("[^a-zA-Z]", " ").split(" "))
    sentences.pop() 
    
    return sentences

def sentence_similarity(sent1, sent2, stopwords=None):
    if stopwords is None:
        stopwords = []
 
    sent1 = [w.lower() for w in sent1]
    sent2 = [w.lower() for w in sent2]
 
    all_words = list(set(sent1 + sent2))
 
    vector1 = [0] * len(all_words)
    vector2 = [0] * len(all_words)
 
    for w in sent1:
        if w in stopwords:
            continue
        vector1[all_words.index(w)] += 1
 
    for w in sent2:
        if w in stopwords:
            continue
        vector2[all_words.index(w)] += 1
 
    return 1 - cosine_distance(vector1, vector2)
 
def build_similarity_matrix(sentences, stop_words):
    similarity_matrix = np.zeros((len(sentences), len(sentences)))
 
    for idx1 in range(len(sentences)):
        for idx2 in range(len(sentences)):
            if idx1 == idx2: 
                continue 
            similarity_matrix[idx1][idx2] = sentence_similarity(sentences[idx1], sentences[idx2], stop_words)

    return similarity_matrix


def generate_summary(file_name, top_n=5):
    stop_words = stopwords.words('english')
    summarize_text = []

    sentences =  read_article(file_name)

    sentence_similarity_martix = build_similarity_matrix(sentences, stop_words)

    sentence_similarity_graph = nx.from_numpy_array(sentence_similarity_martix)
    scores = nx.pagerank(sentence_similarity_graph)

    ranked_sentence = sorted(((scores[i],s) for i,s in enumerate(sentences)), reverse=True)    

    for i in range(top_n):
      summarize_text.append(" ".join(ranked_sentence[i][1]))
    final = ''
    for i in summarize_text:
      final += i + '.\n'
    return final

def valid_xml_char_ordinal(c):
    codepoint = ord(c)
    # conditions ordered by presumed frequency
    return (
        0x20 <= codepoint <= 0xD7FF or
        codepoint in (0x9, 0xA, 0xD) or
        0xE000 <= codepoint <= 0xFFFD or
        0x10000 <= codepoint <= 0x10FFFF
        )

def getParams(x, s, e):
  pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe"
  images = convert_from_path(str(x.filename), poppler_path=r'C:\Users\owner\Desktop\Visual Studio Code Projects\release\poppler-0.90.1\bin')
  a = []
  a.append(int(s.get()))
  a.append(int(e.get()))
  summary = ''
  for i in range(a[0]-1, a[1]):
    summary += pytesseract.image_to_string(images[i])
  summary = summary.replace('\n', '')
  summary = ''.join(c for c in summary if valid_xml_char_ordinal(c))
  sentCount = len(sent_tokenize(summary))
  file = open('notes.txt', 'w')    
  file.write(summary)
  file.close()
  summary = generate_summary('notes.txt', int(sentCount//3))
  summary = summary.replace('.', '.\n')
  
  document = Document()
  document.add_heading('Notes', 0)
  summary = sent_tokenize(summary)
  for i in summary:
    p = document.add_paragraph(str(i))
  document.save('notes'+'.docx')
  os.startfile('notes'+'.docx')

def uploadFunc(x):
  x.filename = filedialog.askopenfilename()
  startTitle = Label(x, text='What page would you like to begin on?')
  start = Entry(x)
  endTitle = Label(x, text="What page would you like to end on?")
  end = Entry(x)
  startTitle.pack()
  start.pack()
  endTitle.pack()
  end.pack()
  submitParams = Button(x, text='Submit', bg='white', fg='black', command=lambda:getParams(x, start, end))
  submitParams.pack()

count2 = 0
def smartNotesCommand():
  global count2
  count2 += 1
  if count2 == 1:
    windowNotes = Toplevel(root)
    windowNotes.title('SmartNotes')
    notesButtonFont = tkFont.Font(family='Arial', size=15)
    spacer5 = Label(windowNotes, text='', font=notesButtonFont)
    spacer5.pack()
    upload = Button(windowNotes, text='Upload PDF', bg='white', fg='black', font=notesButtonFont, command=lambda:uploadFunc(windowNotes))
    spacer6 = Label(windowNotes, text='', font=notesButtonFont)
    upload.pack()
    summary = Label(windowNotes, text='', font=notesButtonFont)
    spacer6.pack()

noteIm = Image.open('smartnotes.png')
noteIm = noteIm.resize((270, 211))
noteIm = ImageTk.PhotoImage(noteIm)   
smartNotesButton = Button(root, image=noteIm, borderwidth=0, bg='white', fg='black', font=ButtonFont, command=smartNotesCommand)
smartNotesButton.pack()

#Spacer
spacer1 = Label(root, text='', font=ButtonFont)
spacer1.pack()

root.mainloop()